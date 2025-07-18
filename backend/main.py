from flask import Flask, request, jsonify, send_from_directory
import os
import tempfile
import shutil
from pathlib import Path
from werkzeug.utils import secure_filename
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
import easyocr
from flask_cors import CORS
import fitz  # PyMuPDF
from docx import Document
import importlib.util
from PIL import Image
import pytesseract

app = Flask(__name__)
CORS(app)

# Persistent downloads folder
DOWNLOAD_DIR = os.path.abspath("downloads")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# Model paths (adjust accordingly)
base_paths = {
    "translation": "../Helsinkifr",
    "question_generation": "../potsawee",
    "question_answering": "../Deepset",
    "summarization": "../Distilbard"
}

# Load models safely
models = {}

def load_models():
    # Translation
    if os.path.exists(base_paths["translation"]):
        print("✅ Translation model found.")
        models["translation"] = {
            "tokenizer": AutoTokenizer.from_pretrained(base_paths["translation"]),
            "model": AutoModelForSeq2SeqLM.from_pretrained(base_paths["translation"])
        }
    else:
        print("❌ Translation model not found at:", base_paths["translation"])
        models["translation"] = None

    # Question Generation
    if os.path.exists(base_paths["question_generation"]):
        print("✅ Question generation model found.")
        models["question_generation"] = {
            "tokenizer": AutoTokenizer.from_pretrained(base_paths["question_generation"]),
            "model": AutoModelForSeq2SeqLM.from_pretrained(base_paths["question_generation"])
        }
    else:
        print("❌ Question generation model not found at:", base_paths["question_generation"])
        models["question_generation"] = None

    # Question Answering
    if os.path.exists(base_paths["question_answering"]):
        print("✅ QA model found.")
        models["question_answering"] = pipeline(
            "question-answering",
            model=base_paths["question_answering"],
            tokenizer=base_paths["question_answering"]
        )
    else:
        print("❌ QA model not found at:", base_paths["question_answering"])
        models["question_answering"] = None

    # Summarization
    if os.path.exists(base_paths["summarization"]):
        print("✅ Summarization model found.")
        models["summarization"] = {
            "tokenizer": AutoTokenizer.from_pretrained(base_paths["summarization"]),
            "model": AutoModelForSeq2SeqLM.from_pretrained(base_paths["summarization"])
        }
    else:
        print("❌ Summarization model not found at:", base_paths["summarization"])
        models["summarization"] = None

load_models()

# EasyOCR
reader = easyocr.Reader(['fr'], gpu=True)

# Tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe" # Adjust path as needed
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

# Rule-based module
rule_path = os.path.join(os.path.dirname(__file__), "RULE_BASED_backend.py")
spec = importlib.util.spec_from_file_location("rule_module", rule_path)
rule_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(rule_module)

# OCR
def extract_raw_text(image_path, engine="easyocr"):
    if engine == "tesseract":
        image = Image.open(image_path)
        return pytesseract.image_to_string(image, lang="fra")
    else:
        results = reader.readtext(str(image_path), detail=0)
        return " ".join(results)

# AI functions
def translate_french_to_english(text):
    model_info = models["translation"]
    if model_info is None:
        return "Translation model not found at specified path."
    tokenizer = model_info["tokenizer"]
    model = model_info["model"]
    inputs = tokenizer(text, return_tensors="pt").to(model.device)
    outputs = model.generate(**inputs)
    return tokenizer.decode(outputs[0], skip_special_tokens=True)

def generate_questions(text):
    model_info = models["question_generation"]
    if model_info is None:
        return ["Question generation model not found at specified path."]
    tokenizer = model_info["tokenizer"]
    model = model_info["model"]
    prompt = f"generate questions: {text}"
    inputs = tokenizer(prompt, return_tensors="pt").to(model.device)
    outputs = model.generate(**inputs, max_length=128, num_return_sequences=3, do_sample=True)
    return [tokenizer.decode(o, skip_special_tokens=True).split('?')[0].strip() + '?' for o in outputs]

def answer_questions(context, questions):
    qa_model = models["question_answering"]
    if qa_model is None:
        return [(q, "QA model not found at specified path.") for q in questions]
    return [(q, qa_model(question=q, context=context)["answer"]) for q in questions]

def summarize_text(text):
    model_info = models["summarization"]
    if model_info is None:
        return "Summarization model not found at specified path."
    tokenizer = model_info["tokenizer"]
    model = model_info["model"]
    inputs = tokenizer(text, return_tensors="pt").to(model.device)
    outputs = model.generate(**inputs, max_length=100, min_length=30, length_penalty=2.0)
    return tokenizer.decode(outputs[0], skip_special_tokens=True)

# PDF to images
def pdf_to_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        image_path = os.path.join(tempfile.gettempdir(), f"page_{page_num}.png")
        pix.save(image_path)
        images.append(image_path)
    pdf_document.close()
    return images

# Processing logic
def process_single_page_ai(raw_text, mode):
    if mode == "extract":
        return raw_text
    elif mode == "translate":
        return translate_french_to_english(raw_text)
    elif mode == "summarize":
        return summarize_text(raw_text)
    elif mode == "quiz":
        questions = generate_questions(raw_text)
        return answer_questions(raw_text, questions)

def process_single_page_rule(raw_text, mode):
    if mode == "extract":
        return raw_text
    elif mode == "translate":
        return rule_module.translate(raw_text)
    elif mode == "summarize":
        return rule_module.summarize(raw_text)
    elif mode == "quiz":
        return rule_module.quizify(raw_text, n=3)

# Run pipeline
def run_pipeline(pdf_path, mode, process_method):
    image_files = pdf_to_images(pdf_path)
    processed_pages = []

    for img_path in image_files:
        if process_method == "rule":
            raw_text = extract_raw_text(img_path, engine="tesseract")
            output = process_single_page_rule(raw_text, mode)
        else:
            raw_text = extract_raw_text(img_path)  # EasyOCR
            output = process_single_page_ai(raw_text, mode)

        processed_pages.append({
            "raw_text": raw_text,
            "output": output
        })
        os.remove(img_path)

    return processed_pages

# Create outputs
def create_download_files(processed_pages, mode, base_filename):
    txt_path = os.path.join(DOWNLOAD_DIR, f"{base_filename}.txt")
    docx_path = os.path.join(DOWNLOAD_DIR, f"{base_filename}.docx")

    with open(txt_path, 'w', encoding='utf-8') as f:
        for i, page in enumerate(processed_pages):
            f.write(f"=== Page {i+1} ===\n\n")
            if mode == "quiz" and isinstance(page["output"], list):
                for q, a in page["output"]:
                    f.write(f"Q: {q}\nA: {a}\n\n")
            else:
                f.write(f"{page['output']}\n\n")
            f.write("=" * 40 + "\n\n")

    doc = Document()
    doc.add_heading(f'Document Processing: {mode.capitalize()}', 0)
    for i, page in enumerate(processed_pages):
        doc.add_heading(f'Page {i+1}', level=1)
        if mode == "quiz" and isinstance(page["output"], list):
            for q, a in page["output"]:
                doc.add_paragraph(f"Q: {q}", style='ListBullet')
                doc.add_paragraph(f"A: {a}")
        else:
            doc.add_paragraph(str(page['output']))
        if i < len(processed_pages) - 1:
            doc.add_page_break()
    doc.save(docx_path)

    return txt_path, docx_path

# API Endpoint
@app.route('/<action>', methods=['POST'])
def handle_action(action):
    if 'pdf' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['pdf']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Only PDF files allowed'}), 400

    allowed_actions = ["extract", "translate", "summarize", "quiz"]
    if action not in allowed_actions:
        return jsonify({'error': f'Invalid action. Allowed: {allowed_actions}'}), 400

    process_method = request.form.get("process_method", "ai").lower()
    if process_method not in ["ai", "rule"]:
        process_method = "ai"

    temp_dir = tempfile.mkdtemp()
    try:
        pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(pdf_path)

        processed_pages = run_pipeline(pdf_path, action, process_method)
        base_filename = Path(file.filename).stem + f"_{action}_{process_method}"

        txt_path, docx_path = create_download_files(processed_pages, action, base_filename)

        return jsonify({
            "page1": processed_pages[0]["output"],
            "total_pages": len(processed_pages),
            "download_txt": f"/downloads/{Path(txt_path).name}",
            "download_docx": f"/downloads/{Path(docx_path).name}",
            "filename": base_filename
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Download route
@app.route('/downloads/<filename>')
def download(filename):
    return send_from_directory(DOWNLOAD_DIR, filename, as_attachment=True)

# Run the app
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
