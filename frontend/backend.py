from flask import Flask, request, jsonify, send_from_directory
import os
import tempfile
import shutil
from pathlib import Path
from werkzeug.utils import secure_filename
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
import easyocr
from flask_cors import CORS
import fitz  # PyMuPDF
from docx import Document
import importlib.util

app = Flask(__name__)
CORS(app)

# Create persistent downloads folder
DOWNLOAD_DIR = os.path.abspath("downloads")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# Model paths (unchanged)
base_paths = {
    "translation": "../Helsinkifr",
    "question_generation": "../potsawee",
    "question_answering": "../Deepset",
    "summarization": "../Distilbard"
}

# Load models (unchanged)
print("🔄 Loading AI models...")
trans_tokenizer = AutoTokenizer.from_pretrained(base_paths["translation"])
trans_model = AutoModelForSeq2SeqLM.from_pretrained(base_paths["translation"])
qg_tokenizer = AutoTokenizer.from_pretrained(base_paths["question_generation"])
qg_model = AutoModelForSeq2SeqLM.from_pretrained(base_paths["question_generation"])
qa_pipeline = pipeline("question-answering", 
                      model=base_paths["question_answering"],
                      tokenizer=base_paths["question_answering"])
sum_tokenizer = AutoTokenizer.from_pretrained(base_paths["summarization"])
sum_model = AutoModelForSeq2SeqLM.from_pretrained(base_paths["summarization"])
reader = easyocr.Reader(['fr'], gpu=True)

# --- RULE-BASED IMPORTS ---
# Dynamically import rule.py as rule_module
rule_path = os.path.join(os.path.dirname(__file__), "rule.py")
spec = importlib.util.spec_from_file_location("rule_module", rule_path)
rule_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(rule_module)

# Utility functions (unchanged)
def extract_raw_text(image_path):
    results = reader.readtext(str(image_path), detail=0)
    return " ".join(results)

def translate_french_to_english(text):
    inputs = trans_tokenizer(text, return_tensors="pt").to(trans_model.device)
    outputs = trans_model.generate(**inputs)
    return trans_tokenizer.decode(outputs[0], skip_special_tokens=True)

def generate_questions(text):
    prompt = f"generate questions: {text}"
    inputs = qg_tokenizer(prompt, return_tensors="pt").to(qg_model.device)
    outputs = qg_model.generate(**inputs, max_length=128, num_return_sequences=3, do_sample=True)
    return [qg_tokenizer.decode(o, skip_special_tokens=True).split('?')[0].strip() + '?' for o in outputs]

def answer_questions(context, questions):
    return [(q, qa_pipeline(question=q, context=context)["answer"]) for q in questions]

def summarize_text(text):
    inputs = sum_tokenizer(text, return_tensors="pt").to(sum_model.device)
    outputs = sum_model.generate(**inputs, max_length=100, min_length=30, length_penalty=2.0)
    return sum_tokenizer.decode(outputs[0], skip_special_tokens=True)

def pdf_to_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        image_path = os.path.join(tempfile.gettempdir(), f"page_{page_num}.png")
        pix.save(image_path)
        images.append(image_path)
    pdf_document.close()
    return images

def process_single_page_ai(raw_text, mode):
    if mode == "extract":
        return raw_text
    elif mode == "translate":
        return translate_french_to_english(raw_text)
    elif mode == "summarize":
        return summarize_text(raw_text)
    elif mode == "quiz":
        questions = generate_questions(raw_text)
        return answer_questions(raw_text, questions)

def process_single_page_rule(raw_text, mode):
    if mode == "extract":
        return raw_text
    elif mode == "translate":
        return rule_module.translate(raw_text)
    elif mode == "summarize":
        return rule_module.summarize(raw_text)
    elif mode == "quiz":
        questions = rule_module.quizify(raw_text, n=3)
        # For rule-based, just return questions with empty answers
        return [(q, "") for q in questions]

def run_pipeline(pdf_path, mode, process_method):
    image_files = pdf_to_images(pdf_path)
    processed_pages = []

    for img_path in image_files:
        raw_text = extract_raw_text(img_path)
        if process_method == "rule":
            output = process_single_page_rule(raw_text, mode)
        else:
            output = process_single_page_ai(raw_text, mode)
        processed_pages.append({
            "raw_text": raw_text,
            "output": output
        })
        os.remove(img_path)

    return processed_pages

def create_download_files(processed_pages, mode, base_filename):
    txt_path = os.path.join(DOWNLOAD_DIR, f"{base_filename}.txt")
    docx_path = os.path.join(DOWNLOAD_DIR, f"{base_filename}.docx")
    
    # Write TXT
    with open(txt_path, 'w', encoding='utf-8') as f:
        for i, page in enumerate(processed_pages):
            f.write(f"=== Page {i+1} ===\n\n")
            if mode == "quiz":
                for q, a in page["output"]:
                    f.write(f"Q: {q}\nA: {a}\n\n")
            else:
                f.write(f"{page['output']}\n\n")
            f.write("="*40 + "\n\n")
    
    # Write DOCX
    doc = Document()
    doc.add_heading(f'Document Processing: {mode.capitalize()}', 0)
    for i, page in enumerate(processed_pages):
        doc.add_heading(f'Page {i+1}', level=1)
        if mode == "quiz":
            for q, a in page["output"]:
                doc.add_paragraph(f"Q: {q}", style='ListBullet')
                doc.add_paragraph(f"A: {a}")
        else:
            doc.add_paragraph(page['output'])
        if i < len(processed_pages) - 1:
            doc.add_page_break()
    doc.save(docx_path)
    
    return txt_path, docx_path

@app.route('/<action>', methods=['POST'])
def handle_action(action):
    if 'pdf' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['pdf']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Only PDF files allowed'}), 400

    allowed_actions = ["extract", "translate", "summarize", "quiz"]
    if action not in allowed_actions:
        return jsonify({'error': f'Invalid action. Allowed: {allowed_actions}'}), 400

    # Get process_method from form, default to 'ai'
    process_method = request.form.get("process_method", "ai").lower()
    if process_method not in ["ai", "rule"]:
        process_method = "ai"

    temp_dir = tempfile.mkdtemp()
    try:
        pdf_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(pdf_path)

        processed_pages = run_pipeline(pdf_path, action, process_method)
        base_filename = Path(file.filename).stem + f"_{action}_{process_method}"

        txt_path, docx_path = create_download_files(processed_pages, action, base_filename)

        return jsonify({
            "page1": processed_pages[0]["output"],
            "total_pages": len(processed_pages),
            "download_txt": f"/downloads/{Path(txt_path).name}",
            "download_docx": f"/downloads/{Path(docx_path).name}",
            "filename": base_filename
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Leave temp_dir removal to separate cleanup process
        pass

@app.route('/downloads/<filename>')
def download(filename):
    return send_from_directory(DOWNLOAD_DIR, filename, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
